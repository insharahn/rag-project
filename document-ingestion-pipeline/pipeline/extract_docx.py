"""Extract text from DOCX documents.

Word files don't have a real concept of "pages" inside the file itself
(pagination depends on how it's rendered), so unlike the PDF extractor
there's no page-by-page breakdown here — just full text pulled from
paragraphs and tables, in document order.
"""

import sys

from docx import Document


def extract_docx_text(path: str) -> dict:
    """Extract text from a DOCX file: paragraphs and table cells.

    Returns a dict with:
      - full_text: all extracted text, paragraphs and tables included
      - paragraph_count: number of non-empty paragraphs found
      - table_count: number of tables found
    """
    doc = Document(path)

    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    table_rows = []
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip(" |"):
                table_rows.append(row_text)

    full_text = "\n\n".join(paragraphs + table_rows)

    return {
        "full_text": full_text,
        "paragraph_count": len(paragraphs),
        "table_count": len(doc.tables),
    }


if __name__ == "__main__":
    if len(sys.argv) != 2:
        print("Usage: python pipeline/extract_docx.py <path-to-docx>")
        sys.exit(1)

    result = extract_docx_text(sys.argv[1])

    print(f"Paragraphs: {result['paragraph_count']}")
    print(f"Tables: {result['table_count']}")
    print("\n--- First 500 characters of extracted text ---\n")
    print(result["full_text"][:500])